from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

questions = [
    ("Sweet cravings", "Salty/savory cravings", "Sweet", "Salty"),
    ("Morning sickness", "No morning sickness", "Morning sickness", "No morning sickness"),
    ("Carrying high", "Carrying low", "High", "Low"),
    ("Glowing skin", "Breaking out", "Glowing", "Breaking out"),
    ("Heart rate above 140", "Heart rate below 140", "Above 140", "Below 140"),
    ("Moody & emotional", "Calm & chill", "Moody", "Calm"),
    ("Hair getting thinner", "Hair getting thicker & fuller", "Thinner", "Thicker"),
    ("Cold feet", "Warm feet", "Warm", "Cold"),
    ("Craving fruit", "Craving meat & cheese", "Fruit", "Meat & cheese"),
    ("Bump out front (like a ball)", "Bump spread wide", "Out front", "Wide"),
    ("Sleeping on your left side", "Sleeping on your right side", "Left", "Right"),
    ("Headaches often", "No headaches", "No headaches", "Headaches"),
    ("Craving orange juice", "Craving apple juice", "Orange juice", "Apple juice"),
    ("Dad gaining weight too", "Dad staying the same", "Dad gaining", "Dad same"),
    ("Dry hands", "Soft hands", "Soft", "Dry"),
    ("Bump low & like a watermelon", "Bump neat & like a basketball", "Watermelon", "Basketball"),
    ("Nails growing fast & strong", "Nails weak & brittle", "Strong nails", "Brittle nails"),
    ("Leg hair growing faster", "Leg hair same as usual", "Same", "Faster"),
    ("More tired than ever", "Energy levels normal", "More tired", "Normal energy"),
    ("Nose spreading/getting wider", "Nose staying the same", "Same nose", "Nose spreading"),
]

doc = Document()

title = doc.add_heading("This or That - Gender Reveal Edition", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

intro = doc.add_paragraph(
    "Tahnee picks an answer for each pair below. Each option points to either "
    "a Girl or a Boy. Tally up the results to reveal the prediction!"
)
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER

table = doc.add_table(rows=1, cols=5)
table.style = "Light Grid Accent 1"

hdr_cells = table.rows[0].cells
headers = ["#", "This", "That", "Girl Answer", "Boy Answer"]
for cell, text in zip(hdr_cells, headers):
    cell.text = text
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

for i, (this, that, girl, boy) in enumerate(questions, start=1):
    row_cells = table.add_row().cells
    row_cells[0].text = str(i)
    row_cells[1].text = this
    row_cells[2].text = that
    row_cells[3].text = girl
    row_cells[4].text = boy

for row in table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

doc.save("/home/user/gender-reveal/This_or_That_Gender_Reveal.docx")
print("done")
